"""
Genera el documento de la practica de Verificacion del Software:
Auditoria de QA sobre el TFG (7 principios ISTQB) - VERSION RECORTADA.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Margenes
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Estilo base
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

NAVY = RGBColor(0x1A, 0x3A, 0x52)
RED = RGBColor(0xC0, 0x39, 0x2B)
GRAY = RGBColor(0x55, 0x55, 0x55)


def h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)


def h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RED
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)


def para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.bold = bold
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)
    p.paragraph_format.left_indent = Cm(0.6)


def code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = GRAY
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(6)


# ============================================================
# PORTADA - el usuario la retoca a mano, dejamos minimalista
# ============================================================
p = doc.add_paragraph()
run = p.add_run("Verificación del Software")
run.font.size = Pt(11)
run.font.color.rgb = GRAY

p = doc.add_paragraph()
run = p.add_run("Auditoría de Aseguramiento de la Calidad (QA)\nsobre el Software del TFG")
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = NAVY
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(16)

para("Aplicación crítica de los 7 principios del ISTQB sobre el sistema desarrollado en el Trabajo de Fin de Grado: \"Ingeniería de variables para la predicción de la Copa Mundial de la FIFA 2026 mediante aprendizaje automático\".")

p = doc.add_paragraph()
run = p.add_run("Autor: ")
run.font.bold = True
p.add_run("Raúl Ramírez Adarve")

p = doc.add_paragraph()
run = p.add_run("Grado: ")
run.font.bold = True
p.add_run("Ingeniería del Software (mención en Ingeniería de Datos)")

p = doc.add_paragraph()
run = p.add_run("Asignatura: ")
run.font.bold = True
p.add_run("Verificación del Software (25/26)")

p = doc.add_paragraph()
run = p.add_run("Universidad: ")
run.font.bold = True
p.add_run("U-tad")

p = doc.add_paragraph()
run = p.add_run("Repositorio: ")
run.font.bold = True
p.add_run("(enlace al repositorio del TFG)")

doc.add_page_break()

# ============================================================
# CONTEXTO BREVE
# ============================================================
h1("Contexto del proyecto")

para("El TFG es un sistema de predicción del Mundial 2026 con tres componentes: un pipeline de machine learning en Python (XGBoost optimizado con Optuna, alimentado por 49.215 partidos internacionales y 50 variables predictivas), un simulador Monte Carlo que ejecuta 10.000 ediciones del torneo, y una aplicación web (FastAPI + Next.js) que sirve los resultados. El código se organiza en model/ (pipeline ML), backend/ (API) y frontend/ (interfaz). A lo largo de esta auditoría se referencian rutas concretas del repositorio.")

# ============================================================
# BLOQUE 1
# ============================================================
h1("Bloque 1. Las pruebas demuestran la presencia de defectos, no su ausencia")

para("Tres defectos reales encontrados durante el desarrollo del TFG:")

h2("Defecto 1. Inflación de ratings ELO en selecciones menores")

para("La primera versión del ELO (model/features/elo.py) usaba K = 40 sin ponderar por importancia de torneo. Al inspeccionar el ranking generado, aparecían selecciones como Islas Feroe o San Marino con ratings cercanos a 1.700, por encima de selecciones objetivamente mejores. La causa: con K = 40, una racha de tres o cuatro amistosos ganados a rivales de su nivel inflaba el rating, porque el sistema no distinguía entre amistoso y partido de Mundial. Se corrigió bajando K a 32, añadiendo pesos por tipo de torneo (de 0,6 para amistosos a 1,5 para Mundial) y una regresión a la media del 0,3 % tras cada actualización.")

h2("Defecto 2. Desalineación entre FEATURE_COLS y el constructor de features")

para("El modelo se entrena con un vector de 50 features cuyo orden está fijado en FEATURE_COLS (model/config.py). El simulador (model/simulation/simulator.py) construye ese vector en el método _build_match_features. Al añadir las tres últimas variables del Ranking FIFA, actualicé FEATURE_COLS pero olvidé añadirlas al numpy.array de _build_match_features. El modelo entrenaba con 50 columnas pero la API devolvía resultados sin sentido porque el vector de inferencia tenía 47 valores y el orden estaba corrido. Es un bug silencioso, sin error en runtime hasta que scikit-learn detectaba el desfase de dimensiones.")

h2("Defecto 3. Fuga temporal en las variables del Ranking FIFA")

para("Las tres features del Ranking FIFA se cargan desde un diccionario estático FIFA_RANKINGS con las posiciones de abril de 2026. Cuando el pipeline procesa partidos del histórico (por ejemplo, uno de 2005), está usando la posición FIFA de 2026 en lugar de la histórica. Es data leakage real, acotado a 3 de las 50 variables. Está documentado como limitación reconocida en la memoria del TFG (sección 5.1).")

h2("Por qué no se puede garantizar que el sistema sea perfecto")

para("Haber corregido estos defectos mejora el sistema, pero no permite afirmar que esté libre de errores. Las pruebas exhaustivas sobre la combinación de 49.215 partidos, 50 variables y 10.000 simulaciones son inviables (Bloque 2). La calidad de las predicciones depende de hipótesis no observables (representatividad del dataset, estabilidad de las relaciones en el tiempo) que ninguna prueba puede validar. Existen defectos silenciosos por construcción: el K = 32 podría ser subóptimo para ciertas confederaciones, la heurística de penaltis es una simplificación. Y por la propia naturaleza del problema, el techo de precisión está en el 55-60 %, por lo que el sistema fallará en una proporción significativa de partidos por diseño. Las pruebas realizadas demuestran la presencia de defectos; ninguna prueba demuestra lo contrario.")

doc.add_page_break()

# ============================================================
# BLOQUE 2
# ============================================================
h1("Bloque 2. Las pruebas exhaustivas son imposibles")

h2("Endpoint elegido: POST /api/predict")

para("Este endpoint es el núcleo de la predicción puntual. Está definido en backend/api/routes.py y acepta un cuerpo JSON con tres parámetros:")

code_block("""class MatchRequest(BaseModel):
    home_team: str
    away_team: str
    is_knockout: bool = False""")

h2("Cálculo de la combinatoria total")

para("Sin acotar el dominio, las combinaciones son: home_team y away_team son strings UTF-8 arbitrarios; aunque limitemos a 100 caracteres ASCII (95 imprimibles), el espacio es 95 elevado a 100, del orden de 5 por 10 elevado a 197 cadenas distintas. Multiplicado por sí mismo (away_team) y por 2 (is_knockout), el total supera cualquier capacidad de cómputo razonable. Ni ejecutando un caso por nanosegundo durante toda la edad del universo se cubriría una fracción significativa.")

para("Acotando al dominio real (aproximadamente 240 nombres de selecciones válidas entre FIFA y selecciones históricas), quedan 240 multiplicado por 240 multiplicado por 2 = 115.200 combinaciones. Sigue siendo inviable: cada predicción tarda 30-50 ms y probarlas todas llevaría más de una hora de ejecución continua.")

h2("Mitigación con particiones de equivalencia y BVA")

para("Identifico cuatro particiones por nombre de equipo y dos para is_knockout: (A) nombre válido en el dataset, (B) nombre válido sintácticamente pero no presente, (C) home y away iguales, (D) cadena vacía o no válida. Combinando particiones y aplicando BVA sobre is_knockout, propongo 5 casos que cubren todas las clases:")

t = doc.add_table(rows=6, cols=3)
t.style = 'Light Grid Accent 1'
hdr = t.rows[0].cells
hdr[0].text = "Caso"
hdr[1].text = "Entrada"
hdr[2].text = "Resultado esperado"

cases = [
    ("1. Partición A + is_knockout=false", "(\"Spain\", \"France\", false)", "200 OK; 3 probabilidades que suman 1,0 coherentes con los ELO actuales."),
    ("2. Partición A + is_knockout=true", "(\"Argentina\", \"Brazil\", true)", "200 OK; activa rama de eliminatoria; probabilidades válidas."),
    ("3. Partición C: equipos iguales", "(\"Spain\", \"Spain\", false)", "Ideal: 400 con mensaje. Actual: empate por construcción del ELO."),
    ("4. Partición B: equipos inexistentes", "(\"Atlantis\", \"Wakanda\", false)", "Ideal: 400. Actual: el sistema asigna ELO inicial y FIFA por defecto, devuelve probabilidades simétricas."),
    ("5. Partición D: cadena vacía", "(\"\", \"France\", false)", "422 de Pydantic antes de invocar al simulador, exigiendo min_length=1 en el esquema."),
]
for i, (tipo, inp, exp) in enumerate(cases, start=1):
    row = t.rows[i].cells
    row[0].text = tipo
    row[1].text = inp
    row[2].text = exp

para(" ")

para("Estos 5 casos no demuestran que la función sea correcta; demuestran que se ha cubierto un representante de cada clase de entrada. Los casos 3, 4 y 5 ponen al descubierto validaciones que actualmente no existen y que conviene añadir: rechazar equipos iguales, exigir que el nombre esté en el dataset y forzar min_length=1 en el esquema Pydantic.")

doc.add_page_break()

# ============================================================
# BLOQUE 3
# ============================================================
h1("Bloque 3. Pruebas tempranas (Early testing / Shift-Left)")

h2("Cuándo empecé a probar")

para("El ciclo de vida del TFG fue iterativo pero claramente sesgado al final. Las semanas 1-3 fueron prototipo del ELO en notebook con validación visual. Las semanas 4-7 implementación del pipeline con prints. Las semanas 8-10 entrenamiento y comparativa de modelos, donde la única \"prueba\" formal fue el split temporal train/val/test. Las semanas 11-13 simulador y aplicación web, con verificación manual lanzando el sistema. Las semanas 14-16 redacción de la memoria, donde aparecieron algunos bugs porque al usar el sistema repetidamente se rompía la inferencia (el Defecto 2 del Bloque 1 apareció aquí). No hay tests unitarios ni de integración escritos en código. La cobertura efectiva es cero.")

h2("Rediseño con Shift-Left")

para("El Shift-Left consiste en mover el testing lo más a la izquierda posible. Aplicado al TFG, propongo dos actividades concretas que habrían evitado dos de los tres bugs documentados:")

para("En la fase de Ingeniería de Requisitos, escribir criterios de aceptación medibles antes de programar. Por ejemplo: \"ninguna selección con menos de 100 partidos en el dataset puede aparecer en el top 30 del ranking ELO\". Este criterio habría detectado en la primera iteración la inflación de ratings de selecciones menores (Defecto 1) en lugar de descubrirla semanas después, cuando ya había contaminado el primer modelo entrenado.")

para("En la fase de Diseño de la Arquitectura, escribir tests de contrato entre módulos. Por ejemplo, un test tan simple como:")

code_block("""def test_simulator_feature_vector_matches_feature_cols():
    simulator = build_dummy_simulator()
    vector = simulator._build_match_features("Spain", "France")
    assert vector.shape[1] == len(FEATURE_COLS), (
        f"Generadas {vector.shape[1]} columnas, "
        f"FEATURE_COLS define {len(FEATURE_COLS)}"
    )""")

para("Habría detectado al instante el Defecto 2 (desalineación entre FEATURE_COLS y _build_match_features). En lugar de horas de inferencias produciendo predicciones absurdas que era difícil identificar como bug de software, el test habría fallado en el primer commit que añadió las features del Ranking FIFA sin actualizar el simulador.")

doc.add_page_break()

# ============================================================
# BLOQUE 4
# ============================================================
h1("Bloque 4. Agrupación de defectos (Defect clustering)")

h2("Módulos del proyecto por tamaño y responsabilidad")

t = doc.add_table(rows=5, cols=3)
t.style = 'Light Grid Accent 1'
hdr = t.rows[0].cells
hdr[0].text = "Módulo"
hdr[1].text = "Líneas"
hdr[2].text = "Criticidad"

modulos = [
    ("model/visualization/plots.py", "1147", "Baja (gráficas)"),
    ("model/simulation/simulator.py", "366", "Alta (motor del torneo)"),
    ("model/features/engineering.py", "268", "Media-alta (construcción de features)"),
    ("model/models/training.py", "213", "Media (script de entrenamiento)"),
]
for i, (mod, lin, crit) in enumerate(modulos, start=1):
    row = t.rows[i].cells
    row[0].text = mod
    row[1].text = lin
    row[2].text = crit

para(" ")

h2("Módulo central: model/simulation/simulator.py")

para("Aunque plots.py es el más largo en líneas, el módulo donde se concentra la lógica crítica es simulator.py. Combina los tres factores que el principio de defect clustering identifica como aglutinadores de defectos:")

bullet("Complejidad ciclomática alta: loops anidados sobre los 12 grupos, ramas por tipo de resultado (victoria local, empate, visitante), ajustes heurísticos de marcadores, resolución de penaltis en eliminatoria y ensamblado del bracket completo (R32, R16, cuartos, semifinales, final).")
bullet("Acoplamiento muy alto: simulator.py importa de cinco módulos del proyecto (config, elo, tournament, form, h2h). Cualquier defecto en cualquiera de esos se propaga aquí.")
bullet("Lógica de negocio crítica: todo el valor del TFG depende del simulador. Las probabilidades agregadas que el usuario consulta en la web salen de ejecutar este módulo 10.000 veces. Un defecto aquí contamina absolutamente todas las cifras de la aplicación.")

h2("Justificación bajo Pareto 80/20")

para("Distribuir el esfuerzo de testing equitativamente entre los 30 ficheros .py del proyecto desperdiciaría tiempo en módulos como plots.py (cuyo peor fallo es una gráfica visualmente extraña) o config.py (solo constantes, sin lógica que rompa). Concentrarlo en simulator.py es eficiente por tres razones: la probabilidad a priori de defecto por línea es mayor (más condicionales y manipulación de probabilidades), las consecuencias de un defecto allí son catastróficas (se propaga a las 10.000 simulaciones y a toda la web), y los tests del simulador validan implícitamente el flujo completo, porque orquesta el modelo, los trackers y la estructura del torneo. Una asignación realista del esfuerzo: 60-70 % sobre simulator.py y elo.py, 20-25 % sobre engineering.py, y el resto entre backend y visualización.")

doc.add_page_break()

# ============================================================
# BLOQUE 5
# ============================================================
h1("Bloque 5. La paradoja del pesticida")

h2("Suite de pruebas actual: 5 verificaciones manuales repetidas")

t = doc.add_table(rows=6, cols=2)
t.style = 'Light Grid Accent 1'
hdr = t.rows[0].cells
hdr[0].text = "Caso de prueba manual"
hdr[1].text = "Criterio de aceptación"

casos = [
    ("Entrenamiento completo (scripts/optimize_model.py).", "Accuracy XGBoost ≥ 55 %; log loss < 0,95."),
    ("Inspección del top 15 del ranking ELO.", "España, Argentina, Francia, Inglaterra y Brasil aparecen en el top 10. Ninguna selección menor sube de forma anómala."),
    ("Simulación completa (GET /api/simulation?seed=42).", "Exactamente 104 partidos: 72 de grupos + 16 R32 + 8 octavos + 4 cuartos + 2 semis + tercer puesto + final."),
    ("Monte Carlo de 10.000 simulaciones.", "Para cada selección, P(pasa grupos) ≥ P(R16) ≥ P(cuartos) ≥ P(semis) ≥ P(final)."),
    ("POST /api/predict con dos equipos.", "200 con 3 probabilidades que suman 1,0 con tolerancia 1e-6."),
]
for i, (caso, crit) in enumerate(casos, start=1):
    row = t.rows[i].cells
    row[0].text = caso
    row[1].text = crit

para(" ")

h2("Refactor hipotético en Fase 2: sustituir XGBoost por Poisson bivariante")

para("Imagino que en Fase 2 sustituyo el modelo XGBoost de clasificación por dos estimadores Poisson que predicen directamente los goles del local y del visitante. Es una línea futura ya mencionada en la memoria del TFG porque permitiría simulaciones con marcadores estadísticamente más rigurosos. El cambio implica reescribir simulate_match para samplear goles de distribuciones Poisson en lugar de samplear una clase de resultado y aplicar la heurística actual.")

h2("Cómo sufre la suite la paradoja del pesticida")

bullet("Caso 1 pierde sentido: el modelo Poisson no devuelve clases, devuelve goles. Habría que sustituir accuracy por error medio en goles o log loss bivariante.")
bullet("Caso 2 sobrevive intacto: el ELO no depende del modelo de salida.")
bullet("Caso 3 sigue siendo válido a nivel de conteo (104 partidos) pero podría enmascarar defectos: si el nuevo simulador genera marcadores irreales como 12-9, el test seguiría en verde porque solo cuenta partidos. Necesitaría un test adicional sobre la distribución de goles.")
bullet("Caso 4 sobrevive y gana importancia: es exactamente el tipo de consistencia agregada donde aparecen inconsistencias tras un refactor.")
bullet("Caso 5 falla a nivel de contrato: el endpoint cambiaría su respuesta de 3 probabilidades a 2 lambdas Poisson. Hay que reescribirlo.")

para("El patrón es el típico de la paradoja: tests que ya no valen (1 y 5), tests que pueden enmascarar defectos nuevos (3) y huecos en la cobertura de los nuevos riesgos.")

h2("Estrategia de actualización")

bullet("Auditar caso por caso qué tests siguen siendo conceptualmente válidos. Mantener 2 y 4 intactos. Reescribir 1, 3 y 5 con métricas y contratos nuevos.")
bullet("Añadir tests específicos para los riesgos del nuevo enfoque (distribución de goles razonable, lambdas positivas, marcadores extremos con baja probabilidad).")
bullet("Mantener una suite de regresión congelando resultados de la versión actual con semilla fija, para detectar regresiones inesperadas.")
bullet("Disciplina: cada bug nuevo encontrado, un test nuevo añadido a la suite. Así la suite crece donde el sistema falla.")
bullet("Revisión periódica: si la suite lleva meses sin que falle nada, plantearse críticamente si sigue cubriendo riesgos reales o se ha vuelto rutina.")

doc.add_page_break()

# ============================================================
# BLOQUE 6
# ============================================================
h1("Bloque 6. Las pruebas dependen del contexto")

h2("Contexto del TFG")

para("Mi TFG es una herramienta de analítica deportiva: aplicación web pública con un pipeline batch de ML que precomputa resultados y una capa web que los sirve. El usuario es un aficionado al fútbol que entra por curiosidad. Los resultados son probabilísticos por naturaleza. Un error en una predicción no tiene consecuencias materiales para el usuario. La latencia importa pero no es crítica. No hay datos personales ni transacciones. No hay normativa sectorial aplicable más allá de RGPD básico.")

h2("Contraste con un caso opuesto: software embebido médico")

para("Imagino el firmware de un marcapasos o un ventilador de UCI. Es un caso radicalmente opuesto en todos los ejes relevantes:")

t = doc.add_table(rows=6, cols=3)
t.style = 'Light Grid Accent 1'
hdr = t.rows[0].cells
hdr[0].text = "Eje"
hdr[1].text = "Mi TFG"
hdr[2].text = "Software médico"

ejes = [
    ("Rendimiento", "Latencia de pocos ms; aceptable que falle puntualmente bajo carga.", "Latencia determinista en ms garantizada siempre. Cualquier retraso puede ser letal."),
    ("Seguridad", "Mínima. API pública de lectura, sin autenticación.", "Auditoría obligatoria: firmware firmado, comunicaciones cifradas, pruebas de penetración."),
    ("Tolerancia a fallos", "Alta. Si falla, el usuario refresca y listo.", "Cero. Redundancia obligatoria, watchdog, modos de seguridad por hardware."),
    ("Normativa", "Ninguna específica. RGPD básico.", "ISO 13485, IEC 62304, ISO 14971, FDA o CE Marking."),
    ("Cobertura objetivo", "No formalizada. En la práctica, baja.", "MC/DC cerca del 100 % en lógica crítica. Trazabilidad requisito-prueba documentada."),
]
for i, (eje, mio, med) in enumerate(ejes, start=1):
    row = t.rows[i].cells
    row[0].text = eje
    row[1].text = mio
    row[2].text = med

para(" ")

para("La estrategia de testing de mi TFG (exploratoria, validación visual, tests manuales repetidos, sin suite automatizada formal) sería negligente en el contexto médico. Al revés, aplicar el rigor del software médico al TFG sería sobreingeniería que no aportaría valor proporcional al esfuerzo. El contexto manda: hablar en abstracto de \"hay que tener un X % de cobertura\" sin precisar el contexto es un consejo vacío.")

doc.add_page_break()

# ============================================================
# BLOQUE 7
# ============================================================
h1("Bloque 7. Falacia de ausencia de errores")

h2("El experimento mental")

para("Supongamos que una suite de tests masiva se ejecuta sobre el TFG y todo sale en verde: endpoints respondiendo, modelo cargado, simulaciones acabando sin excepciones, web pintando sin errores en consola. Y sin embargo, al entregarlo al usuario final, este descubre que no entiende qué significa que España tenga un 21,9 %, que la web no explica por qué unas selecciones aparecen por encima de otras, que cada simulación interactiva da resultados distintos generando sensación de inconsistencia, y que el sistema no aporta valor diferencial frente a una web de casas de apuestas. Todo eso puede pasar perfectamente con los tests en verde. La falacia de la ausencia de errores consiste en confundir \"sin errores de código\" con \"resuelve el problema del usuario\".")

h2("Autocrítica del TFG")

para("Lo que el TFG sí resuelve: cumple los cinco objetivos específicos de la memoria, aporta valor académico real documentando con SHAP qué variables pesan más (algo que ni Groll ni Ogundeji desarrollan con tanto detalle en sus papers), y es reproducible al 100 % con datos y código públicos, lo que sí lo distingue de las casas de apuestas.")

para("Lo que el TFG no resuelve o no resuelve bien: la explicación al usuario final es pobre, porque la web muestra porcentajes y poco más, sin un \"por qué España\" en lenguaje comprensible (el sistema tiene esa información gracias a SHAP, pero no se la enseña al usuario). Le falta sección de metodología accesible que dignifique las cifras frente a otras webs de predicciones. No hay interactividad real más allá de leer: el usuario no puede cambiar grupos, comparar selecciones cara a cara con explicación detallada, ni explorar escenarios. Las predicciones quedarán obsoletas en cuanto cambien factores externos (lesiones, sorpresas de la fase de grupos) y no hay automatización para reentrenar.")

para("Todos estos problemas son del tipo que ningún test automático detectaría. Un test puede comprobar que el endpoint devuelve un JSON correcto, pero no que ese JSON sea útil. Un test puede asegurar que accuracy es 60,7 %, pero no que ese 60,7 % aporte valor real al usuario frente a alternativas. La única forma de comprobar lo segundo es probarlo con usuarios reales, no con tests. El TFG funciona sin errores de código y cumple los objetivos formales, pero hay aspectos que un usuario final podría legítimamente echar en falta. Reconocerlo no resta mérito al trabajo; lo sitúa en su justa medida y deja claro dónde están las líneas de mejora.")

# Guardar
out = "C:/Users/adarv/Escritorio/INSO/INSO4/TFG/world-cup-predictor/outputs/QA-TFG-Raul-Ramirez.docx"
doc.save(out)
print(f"OK: {out}")
